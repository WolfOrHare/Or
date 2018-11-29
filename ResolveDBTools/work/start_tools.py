# # -*- coding: utf-8 -*-
import docx
import docx.document
import sys
import os
from work import connect_db
import time
from conf import settings

# 重新载入字符集编码
reload(sys)
sys.setdefaultencoding('utf-8')
file = docx.Document()

def get_styles():

    styles = file.styles
    for s in styles:
        print(s.name)

def select_table():

    con_db = connect_db.QuerySQL()
    sqlcommand = '''select table_name from information_schema.tables where table_schema='test_loan_db' and table_type='base table' '''
    # conlist = run.connectdb(dbinfo)

    conlist = con_db.connectdb(ssh_host='47.94.40.127',
                               ssh_usr='appuser', ssh_psw='5zt4BRJ8',
                               sql_addr='rdsiuzzzqiuzzzq.mysql.rds.aliyuncs.com',
                               sql_port=3306,
                               sql_usr='sit_user01',
                               sql_psw='j*IHNifVbxCJ',
                               db_name='test_loan_db')
    server = conlist[0]
    con = conlist[1]
    # 查询所有的数据库表名
    # 只返回了包含t_lon内容的表名
    _header = settings.TABLE_INCLUDE
    tablenamelist  = con_db.get_table_name(con, sqlcommand,_header)
    print('所有有效的数据库表名----》%s'%tablenamelist)

    # 生成docx文件
    dir = os.path.dirname(os.path.abspath('.')) + '\docxfile\ '
    now = time.strftime("%Y-%m-%d_%H", time.localtime(time.time()))
    file_dir =dir  + _header + '_' + now + '.docx'


    # 获取一张表的建表脚本
    # table_data = con_db.show_create_table(con, tablenamelist[10])
    # print('获取一张表的建表脚本----》' + table_data)

    # ‘`’分割数据库查询结果，产生一个列表
    # table_list = table_data.split('`',)
    # print(table_list)

    # 获取所有表的建表脚本
    for i in range(len(tablenamelist)):
        table_data = con_db.show_create_table(con,tablenamelist[i])
        table_list = table_data.split('`',)
        print(table_list)

            # save_document(tablename[i],result)
        # 获取表描述
        head_com = head_commet(table_list)

        file.add_heading(tablenamelist[i] , level=1)
        file.add_heading(head_com, level=4)

        # 设置表头
        table = file.add_table(rows=1, cols=5,style="Medium List 1 Accent 1")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'序号'
        hdr_cells[1].text = u'字段名'  # 第一行的第一列,给这行里面添加文字
        hdr_cells[2].text = u'类型'
        hdr_cells[3].text = u'是否允许为空'
        hdr_cells[4].text = u'描述'

        write_document(table_list)
        file.add_page_break()

    file.save(file_dir)

    server.stop()
    con.close()
    return table_data
def head_commet(resultlist):
    try:
        for i in range(len(resultlist)):
            if 'CHARSET=utf8' in resultlist[i]:
                head_com = resultlist[i].split('COMMENT=',1)[1]
                print(head_com)
        return "---"+head_com
    except Exception as e:
        print(e)


def write_document(resultlist):
    index = 0
    try:
        for i in range(len(resultlist)):
            resultlist_child = resultlist[i].split(' ',)
            if 'CHARSET=utf8' in resultlist[i]:
                head_COMMENT = resultlist[i].split('COMMENT=', )
                print(head_COMMENT)
                break

            if i > 2 and i % 2 == 1:
                print(i, resultlist_child)

                table = file.add_table(rows=1, cols=5, style="Medium List 1 Accent 1")
                hdr_cells = table.rows[0].cells

                if index <= len(resultlist):
                    index += 1

                    hdr_cells[0].text = str(index)

                if 'IDX_' not in resultlist[i]:
                    hdr_cells[1].text = resultlist[i]
                    hdr_cells[2].text = resultlist[i+1].split(' ',)[1]
                # 如果包含CHARSET=utf8那么取出描述添加到表名后，作为整表的描述
                elif 'CHARSET=utf8' in resultlist[i]:
                    break
                else:
                    break


            # 偶数内容：取list第三个参数，
            if i > 3 and i % 2 == 0:
                print(i, resultlist_child)

                # 是否为空列三种情况：默认值为空，或not null或null
                if (resultlist_child[2] == 'DEFAULT' and resultlist_child[3] == 'NULL') or resultlist_child[2] == 'NOT':
                    print(i, resultlist_child)
                    hdr_cells[3].text = resultlist_child[2] + ' ' + resultlist_child[3]
                    try:
                        COMMENT_IDX =resultlist_child.index('COMMENT')
                        hdr_cells[4].text = resultlist_child[COMMENT_IDX+1:]

                    except Exception as e:
                        print('如果取不到数据则写入空数据：',e)
                        hdr_cells[4].text =' '

                elif resultlist_child[2] == 'NULL':
                    print(i, resultlist_child)
                    hdr_cells[3].text = resultlist_child[2]
                    hdr_cells[4].text = resultlist_child[5]



                elif resultlist_child[2] not in('DEFAULT','NOT','NULL'):
                    print("-----字段参数异常-----")
                    print(resultlist_child)
                    print(resultlist_child[2])

                elif 'IDX_' in resultlist[i]:
                    break
                # 写在表名字后作为注释


            elif resultlist[3] in resultlist[:2]:
                break


    except Exception as e:
        print(e)


def save_document(tablename,content):
    document = docx.Document()  # 首先这是包的主要接口，这应该是利用的设计模式的一种，用来创建docx文档，里面也可以包含文档路径(d:\\2.docx)

    document.add_heading('Document Title', 0)  # 这里是给文档添加一个标题，0表示 样式为title，1则为忽略，其他则是Heading{level},具体可以去官网查;

    p = document.add_paragraph('A plain paragraph having some ')  # 这里是添加一个段落
    p.add_run('bold').bold = True  # 这里是在这个段落p里文字some后面添加bold字符
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)  # 这里是添加标题1
    document.add_paragraph('Intense quote', style='IntenseQuote')  # 这里是添加段落，style后面则是样式

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'  # 添加段落，样式为unordered list类型
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'  # 添加段落，样式为ordered list数字类型
    )
    table = document.add_table(rows=1, cols=3)  # 添加一个表格，每行三列
    hdr_cells = table.rows[0].cells  # 表格第一行的所含有的所有列数
    hdr_cells[0].text = 'Qty'  # 第一行的第一列,给这行里面添加文字
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'

    #for item in recordset:
     #   row_cells = table.add_row().cells  # 这是在这个表格第一行 (称作最后一行更好) 下面再添加新的一行
     #   row_cells[0].text = str(item.qty)
     #   row_cells[1].text = str(item.id)
     #   row_cells[2].text = item.desc
    document.add_page_break()  # 添加分页符
    document.save('demo.docx')

if __name__ == '__main__':
    # run = get_styles()
    run = select_table()